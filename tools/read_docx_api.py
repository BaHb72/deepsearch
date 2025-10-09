#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""读取35个API接口文档内容"""

from pathlib import Path

import docx


def read_api_document():
    """读取35个API接口详细内容.docx"""
    doc_path = Path("third_party/35API接口详细内容.docx")

    if not doc_path.exists():
        print(f"文档不存在: {doc_path}")
        return

    doc = docx.Document(doc_path)

    print("=== 35个API接口详细内容 ===\n")

    # 读取所有段落
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:
            print(f"{text}")

    # 读取所有表格
    print("\n=== 表格内容 ===\n")
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i+1}:")
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            if any(row_data):  # 只打印非空行
                print(" | ".join(row_data))


if __name__ == "__main__":
    read_api_document()
